"""Markdown to PDF conversion service."""
import io
import re
from io import BytesIO
from pathlib import Path
from typing import Literal

from markdown_it import MarkdownIt
from weasyprint import HTML, CSS
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

ThemeType = Literal["default", "academic", "technical"]

THEMES = {
    "default": {
        "name": "Default",
        "css_file": "default.css",
    },
    "academic": {
        "name": "Academic",
        "css_file": "academic.css",
    },
    "technical": {
        "name": "Technical",
        "css_file": "technical.css",
    },
}


def get_system_font() -> str:
    """获取系统中文字体名称。"""
    # Windows 系统字体
    simhei = Path("C:/Windows/Fonts/simhei.ttf")
    if simhei.exists():
        return "SimHei"

    msyh = Path("C:/Windows/Fonts/msyh.ttc")
    if msyh.exists():
        return "Microsoft YaHei"

    # macOS / Linux 回退
    return "Noto Sans CJK SC"


class MarkdownConverter:
    """Markdown to PDF converter using WeasyPrint."""

    def __init__(self, themes_dir: Path | None = None):
        self.md = MarkdownIt("commonmark", {"html": True})
        self.themes_dir = themes_dir or Path(__file__).parent.parent / "static" / "css" / "themes"
        self.font_name = get_system_font()

    def to_html(self, markdown_text: str) -> str:
        """Convert markdown to HTML."""
        return self.md.render(markdown_text)

    def to_pdf(
        self,
        markdown_text: str,
        theme: ThemeType = "default",
        title: str = "Document",
    ) -> bytes:
        """Convert markdown to PDF bytes."""
        html_content = self.to_html(markdown_text)

        theme_config = THEMES.get(theme, THEMES["default"])
        css_path = self.themes_dir / theme_config["css_file"]

        # 基础 CSS + 中文字体
        base_css = f"""
        @page {{
            size: A4;
            margin: 2cm;
        }}

        @font-face {{
            font-family: 'SimHei';
            src: local('SimHei'), local('Microsoft YaHei'), local('Noto Sans CJK SC');
        }}

        body {{
            font-family: 'SimHei', 'Microsoft YaHei', 'Noto Sans CJK SC', sans-serif;
            font-size: 12pt;
            line-height: 1.8;
            color: #333;
        }}

        h1, h2, h3, h4, h5, h6 {{
            font-family: 'SimHei', 'Microsoft YaHei', 'Noto Sans CJK SC', sans-serif;
            color: #222;
            margin-top: 1.2em;
            margin-bottom: 0.6em;
            font-weight: bold;
        }}

        h1 {{ font-size: 24pt; border-bottom: 2px solid #333; padding-bottom: 0.3em; }}
        h2 {{ font-size: 18pt; border-bottom: 1px solid #666; padding-bottom: 0.2em; }}
        h3 {{ font-size: 14pt; }}

        p {{
            margin: 0.8em 0;
            text-align: justify;
        }}

        ul, ol {{
            padding-left: 2em;
            margin: 0.5em 0;
        }}

        li {{
            margin: 0.3em 0;
        }}

        code {{
            font-family: 'Consolas', 'Courier New', monospace;
            background-color: #f5f5f5;
            padding: 2px 6px;
            border-radius: 3px;
            font-size: 0.9em;
        }}

        pre {{
            background-color: #f5f5f5;
            padding: 12px;
            border-radius: 5px;
            overflow-x: auto;
            border: 1px solid #ddd;
        }}

        pre code {{
            background: none;
            padding: 0;
        }}

        blockquote {{
            border-left: 4px solid #ddd;
            padding-left: 1em;
            margin: 1em 0;
            color: #666;
        }}

        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 1em 0;
        }}

        th, td {{
            border: 1px solid #ddd;
            padding: 8px 12px;
            text-align: left;
        }}

        th {{
            background-color: #f5f5f5;
            font-weight: bold;
        }}

        a {{
            color: #0066cc;
            text-decoration: none;
        }}

        a:hover {{
            text-decoration: underline;
        }}

        hr {{
            border: none;
            border-top: 1px solid #ddd;
            margin: 2em 0;
        }}
        """

        # 读取主题 CSS（如果存在）
        theme_css = ""
        if css_path.exists():
            theme_css = css_path.read_text(encoding="utf-8")

        # 构建完整 HTML
        full_html = self._build_html_document(html_content, title)

        # 生成 PDF
        html_doc = HTML(string=full_html, base_url=str(self.themes_dir))
        css_doc = CSS(string=base_css + theme_css)

        pdf_buffer = io.BytesIO()
        html_doc.write_pdf(pdf_buffer, stylesheets=[css_doc])

        return pdf_buffer.getvalue()

    def to_word(self, markdown_text: str, title: str = "Document") -> bytes:
        """Convert markdown to Word document bytes."""
        doc = Document()

        # 设置默认字体
        style = doc.styles['Normal']
        style.font.name = self.font_name
        style.font.size = Pt(12)

        # 解析 markdown 并转换为 Word
        lines = markdown_text.split('\n')
        i = 0

        while i < len(lines):
            line = lines[i].strip()

            # 空行
            if not line:
                i += 1
                continue

            # 标题
            if line.startswith('#'):
                level = len(re.match(r'^#+', line).group())
                text = line.lstrip('#').strip()

                if level == 1:
                    p = doc.add_heading(text, level=0)
                elif level <= 6:
                    p = doc.add_heading(text, level=level)
                i += 1
                continue

            # 无序列表
            if line.startswith(('- ', '* ', '+ ')):
                text = line[2:].strip()
                text = self._process_inline_formatting(doc, doc.add_paragraph(style='List Bullet'), text)
                i += 1
                continue

            # 有序列表
            if re.match(r'^\d+\.\s', line):
                text = re.sub(r'^\d+\.\s', '', line)
                text = self._process_inline_formatting(doc, doc.add_paragraph(style='List Number'), text)
                i += 1
                continue

            # 代码块
            if line.startswith('```'):
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                i += 1  # 跳过结束的 ```

                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph(code_text)
                p.style = doc.styles['Normal']
                for run in p.runs:
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
                continue

            # 引用
            if line.startswith('>'):
                text = line.lstrip('>').strip()
                text = self._process_inline_formatting(doc, doc.add_paragraph(style='Quote'), text)
                i += 1
                continue

            # 分隔线
            if re.match(r'^[-*_]{3,}$', line):
                doc.add_paragraph('─' * 50)
                i += 1
                continue

            # 表格 (简单支持)
            if '|' in line and i + 1 < len(lines) and re.match(r'^[\|\-\s]+$', lines[i + 1]):
                table_lines = [line]
                i += 1
                # 跳过分隔行
                i += 1
                # 收集表格行
                while i < len(lines) and '|' in lines[i]:
                    table_lines.append(lines[i])
                    i += 1

                # 解析表格
                rows = [[cell.strip() for cell in row.split('|') if cell.strip()] for row in table_lines]
                if rows:
                    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                    table.style = 'Table Grid'
                    for row_idx, row_data in enumerate(rows):
                        for col_idx, cell_text in enumerate(row_data):
                            cell = table.rows[row_idx].cells[col_idx]
                            cell.text = cell_text
                            if row_idx == 0:  # 表头加粗
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True
                continue

            # 普通段落
            self._process_inline_formatting(doc, doc.add_paragraph(), line)
            i += 1

        # 保存到字节流
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def _process_inline_formatting(self, doc, paragraph, text: str) -> str:
        """处理行内格式（粗体、斜体、代码、链接）。"""
        # 简单的行内代码处理
        text = re.sub(r'`([^`]+)`', r'\1', text)

        # 移除链接语法，保留文本
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)

        # 处理粗体和斜体
        parts = []
        remaining = text

        while remaining:
            # 粗体
            match = re.match(r'^(.*?)\*\*(.+?)\*\*(.*)', remaining, re.DOTALL)
            if match:
                if match.group(1):
                    run = paragraph.add_run(match.group(1))
                    run.font.name = self.font_name
                run = paragraph.add_run(match.group(2))
                run.bold = True
                run.font.name = self.font_name
                remaining = match.group(3)
                continue

            # 斜体
            match = re.match(r'^(.*?)\*(.+?)\*(.*)', remaining, re.DOTALL)
            if match:
                if match.group(1):
                    run = paragraph.add_run(match.group(1))
                    run.font.name = self.font_name
                run = paragraph.add_run(match.group(2))
                run.italic = True
                run.font.name = self.font_name
                remaining = match.group(3)
                continue

            # 没有更多格式
            if remaining:
                run = paragraph.add_run(remaining)
                run.font.name = self.font_name
            break

        return text

    def _build_html_document(self, content: str, title: str) -> str:
        """Build complete HTML document."""
        return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <title>{title}</title>
</head>
<body>
    <div class="document">
        {content}
    </div>
</body>
</html>"""
